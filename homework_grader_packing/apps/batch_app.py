# --- ensure project root on sys.path (robust for streamlit/pyinstaller) ---
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]  # .../homework_grader
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
# -------------------------------------------------------------------------

import io
import json
from datetime import datetime
from pathlib import Path

import streamlit as st
from PIL import Image

# Core logic
from core.grader_core import (
    AppConfig,
    preprocess_for_ocr,
    batch_grade_images,
    ensure_wrongbook_file,
    export_wrongbook_markdown,
)

# Optional: DOCX export (python-docx)
try:
    from docx import Document
    from docx.shared import Pt

    DOCX_OK = True
except Exception:
    DOCX_OK = False


# -------------------------
# Helpers
# -------------------------
def load_config_from_secrets_or_env() -> AppConfig:
    """
    Priority:
      1) st.secrets
      2) config.toml (project root)
      3) OS env (optionally loaded from .env)
      4) defaults
    Note:
      Sensitive fields should be provided via env/.env, not toml.
    """
    import os
    import sys as _sys
    from pathlib import Path

    # 0) Optional: load .env if present (safe no-op if python-dotenv not installed)
    # PATCH: packaged EXE's CWD may not be exe dir, so search explicitly.
    try:
        from dotenv import load_dotenv  # type: ignore

        exe_dir = Path(_sys.executable).resolve().parent  # dist\HomeworkGrader
        candidates = [
            exe_dir / ".env",   # ✅ packaged app next to HomeworkGrader.exe
            ROOT / ".env",      # ✅ dev: project root
            Path.cwd() / ".env" # fallback
        ]
        dotenv_path = next((p for p in candidates if p.exists()), None)
        if dotenv_path:
            load_dotenv(dotenv_path=dotenv_path, override=True)
            print(f"[config] loaded .env from: {dotenv_path}")
        else:
            print(f"[config] .env not found. tried: {[str(p) for p in candidates]}")
    except Exception:
        pass

    # 1) load config.toml if exists
    def _load_toml() -> dict:
        # PATCH: always read from project root, not cwd
        cfg_path = ROOT / "config.toml"
        if not cfg_path.exists():
            return {}
        txt = cfg_path.read_text(encoding="utf-8", errors="ignore")
        # py3.11+: tomllib; fallback to toml
        try:
            import tomllib  # type: ignore
            return tomllib.loads(txt)
        except Exception:
            try:
                import toml  # type: ignore
                return toml.loads(txt)
            except Exception:
                return {}

    toml_cfg = _load_toml()

    def _get_secret(k: str) -> str | None:
        if hasattr(st, "secrets") and k in st.secrets:
            v = str(st.secrets[k])
            return v if v != "" else None
        return None

    def _get_toml(k: str) -> str | None:
        v = toml_cfg.get(k, None)
        if v is None:
            return None
        v = str(v)
        return v if v != "" else None

    def _get_env(k: str) -> str | None:
        v = os.environ.get(k, "")
        return v if v != "" else None

    def pick(k: str, default: str = "") -> str:
        # secrets > toml > env > default
        for getter in (_get_secret, _get_toml, _get_env):
            v = getter(k)
            if v is not None:
                return v
        return default

    cfg = AppConfig(
        provider=pick("PROVIDER", "spark_http"),
        spark_http_url=pick("SPARK_HTTP_URL", ""),
        spark_api_password=pick("SPARK_API_PASSWORD", ""),  # recommend env/.env
        spark_model=pick("SPARK_MODEL", "x1"),
        ocr_url=pick("OCR_URL", ""),
        ocr_app_id=pick("OCR_APP_ID", ""),
        ocr_api_key=pick("OCR_API_KEY", ""),                # recommend env/.env
        ocr_api_secret=pick("OCR_API_SECRET", ""),          # recommend env/.env
    )
    return cfg


def config_is_ok(cfg: AppConfig) -> bool:
    if cfg.provider == "mock":
        return True
    if cfg.provider in ("cloud_ocr_only", "spark_http"):
        if not all([cfg.ocr_url, cfg.ocr_app_id, cfg.ocr_api_key, cfg.ocr_api_secret]):
            return False
    if cfg.provider == "spark_http":
        if not (cfg.spark_http_url and cfg.spark_api_password):
            return False
    return True


def export_wrongbook_docx(jsonl_path: Path, docx_path: Path) -> int:
    if not DOCX_OK:
        raise RuntimeError("python-docx 未安装，无法导出 docx。")

    rows = []
    if jsonl_path.exists():
        for line in jsonl_path.read_text(encoding="utf-8", errors="ignore").splitlines():
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(line)
                if isinstance(obj, dict):
                    rows.append(obj)
            except Exception:
                pass

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("错题本导出", level=1)
    doc.add_paragraph(f"来源：{str(jsonl_path)}")
    doc.add_paragraph(f"导出时间：{datetime.now().isoformat(timespec='seconds')}")
    doc.add_paragraph(f"记录条数：{len(rows)}")

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    def qkey(r):
        q = str(r.get("qid", ""))
        try:
            return (0, int(q))
        except Exception:
            return (1, q)

    rows = sorted(rows, key=qkey)
    for r in rows:
        qid = r.get("qid", "")
        subj = r.get("subject", "")
        verdict = r.get("verdict", "")
        uncertain = bool(r.get("uncertain", False))
        rec = (r.get("recognized_answer") or "").strip()
        comment = (r.get("comment") or "").strip()

        doc.add_heading(f"题 {qid}（{subj}）", level=2)
        doc.add_paragraph(f"判定：{verdict}{'（不确定）' if uncertain else ''}")
        if rec:
            doc.add_paragraph(f"学生答案：{rec}")
        if comment:
            doc.add_paragraph(f"评语：{comment}")

    doc.save(str(docx_path))
    return len(rows)


# -------------------------
# State
# -------------------------
# phase: idle | running | done
if "phase" not in st.session_state:
    st.session_state["phase"] = "idle"
if "job_started" not in st.session_state:
    st.session_state["job_started"] = False
if "batch_run" not in st.session_state:
    st.session_state["batch_run"] = None


def _start_batch_clicked():
    # 点击按钮那一刻就把 phase 置为 running（下一帧 UI 立刻禁用控件）
    st.session_state["phase"] = "running"
    st.session_state["batch_run"] = None
    st.session_state["job_started"] = False


def main():
    # -------------------------
    # UI
    # -------------------------
    st.set_page_config(page_title="作业批改助手（批量）", layout="wide")
    st.title("作业批改助手-批量版v0.3.5")

    disabled = st.session_state["phase"] == "running"

    if st.session_state["phase"] == "running":
        st.warning("批改进行中：界面已锁定，请等待完成…")
    elif st.session_state["phase"] == "done":
        st.success("批改完成 ✅ 你可以开始新一轮。")

    cfg0 = load_config_from_secrets_or_env()

    with st.sidebar:
        st.header("设置")

        provider = st.selectbox(
            "运行模式",
            ["spark_http", "cloud_ocr_only", "mock"],
            index=["spark_http", "cloud_ocr_only", "mock"].index(cfg0.provider)
            if cfg0.provider in ["spark_http", "cloud_ocr_only", "mock"]
            else 0,
            key="provider",
            disabled=disabled,
        )

        subject = st.selectbox("科目", ["科学", "英语"], key="subject", disabled=disabled)
        mode = st.selectbox("批改模式", ["check_and_comment", "check_only"], key="mode", disabled=disabled)

        st.divider()
        st.subheader("题目信息（可选）")
        question_text = st.text_area("题目补充说明（可留空）", value="", height=80, key="question_text", disabled=disabled)
        answer_key = st.text_input("标准答案（可留空）", value="", key="answer_key", disabled=disabled)

        st.divider()
        st.subheader("错题本（按学生分开）")
        enable_wrongbook = st.checkbox("自动保存错题记录", value=True, key="enable_wrongbook", disabled=disabled)
        auto_export_wrongbook = st.checkbox(
            "批量完成后自动导出错题本文档（md/docx）",
            value=False,
            key="auto_export_wrongbook",
            disabled=disabled,
        )

        st.divider()
        st.subheader("图片预处理")
        crop_ratio = st.slider(
            "中心裁剪比例（越小越聚焦）",
            min_value=0.60,
            max_value=1.00,
            value=1.00,
            step=0.02,
            key="crop_ratio",
            disabled=disabled,
        )
        max_side = st.selectbox("最大边长(px)", [1200, 1600, 2000], index=1, key="max_side", disabled=disabled)
        jpeg_quality = st.selectbox("JPEG质量", [60, 70, 75, 80, 85], index=2, key="jpeg_quality", disabled=disabled)
        enable_cleanup = st.checkbox("用大模型清洗OCR文本", value=False, key="enable_cleanup", disabled=disabled)

        st.divider()
        st.subheader("并发（LLM限制=2）")
        llm_concurrency = st.selectbox("LLM并发", [1, 2], index=1, key="llm_concurrency", disabled=disabled)

    cfg = AppConfig(
        provider=provider,
        spark_http_url=cfg0.spark_http_url,
        spark_api_password=cfg0.spark_api_password,
        spark_model=cfg0.spark_model,
        ocr_url=cfg0.ocr_url,
        ocr_app_id=cfg0.ocr_app_id,
        ocr_api_key=cfg0.ocr_api_key,
        ocr_api_secret=cfg0.ocr_api_secret,
    )

    if provider != "mock" and not config_is_ok(cfg):
        st.error("配置不完整：请检查 st.secrets 或环境变量中的 OCR_* / SPARK_*。")
        st.stop()

    col1, col2 = st.columns([1, 1], gap="large")

    with col1:
        st.subheader("上传作业图片（多选）")
        uploaded_files = st.file_uploader(
            "支持 jpg / png（可一次上传 ~50 份）",
            type=["jpg", "jpeg", "png", "webp"],
            accept_multiple_files=True,
            key="batch_uploader",
            disabled=disabled,
        )

        if not uploaded_files:
            st.session_state.pop("preview_file", None)
            st.info("请上传多张作业照片或扫描件，然后点击“开始批改（批量）”。")
            st.stop()

        names = [f.name for f in uploaded_files]
        file_map = {f.name: f for f in uploaded_files}

        old = st.session_state.get("preview_file")
        if old not in names:
            st.session_state["preview_file"] = names[0]

        preview_name = st.selectbox("预览文件", names, key="preview_file", disabled=disabled)
        preview_file = file_map.get(preview_name)
        if preview_file is None:
            st.warning("预览文件已被移除，请重新选择。")
            st.stop()

        raw_bytes = preview_file.getvalue()
        raw_img = Image.open(io.BytesIO(raw_bytes)).convert("RGB")
        st.image(raw_img, caption="原图（预览）", use_container_width=True)

        st.divider()

        st.button(
            "开始批改（批量）",
            type="primary",
            use_container_width=True,
            key="start_batch",
            disabled=disabled or (not uploaded_files),
            on_click=_start_batch_clicked,
        )

    with col2:
        st.subheader("结果")
        if st.session_state["phase"] == "running" and not st.session_state["job_started"]:
            st.session_state["job_started"] = True
            try:
                images = [(f.name, f.getvalue()) for f in uploaded_files]
                with st.status("批改中…（批量）", expanded=True) as status:
                    status.write("已锁定界面控件，正在处理图片…")

                    out = batch_grade_images(
                        cfg=cfg,
                        images=images,
                        subject=subject,
                        mode=mode,
                        question_text=question_text,
                        answer_key=answer_key,
                        enable_cleanup=enable_cleanup,
                        run_root="./runs",
                        max_concurrency=int(llm_concurrency),
                        crop_ratio=float(crop_ratio),
                        max_side=int(max_side),
                        jpeg_quality=int(jpeg_quality),
                        save_wrongbook=bool(enable_wrongbook),
                    )
                    status.update(label="批改完成 ✅", state="complete", expanded=False)

                st.session_state["batch_run"] = out

                if enable_wrongbook and auto_export_wrongbook:
                    for s in out.get("students", []) or []:
                        student_dir = Path(s.get("student_dir", ""))
                        wb = student_dir / "wrongbook.jsonl"
                        ensure_wrongbook_file(str(wb))
                        export_dir = student_dir / "export"
                        export_dir.mkdir(parents=True, exist_ok=True)
                        export_wrongbook_markdown(str(wb), str(export_dir / "wrongbook.md"))
                        if DOCX_OK:
                            export_wrongbook_docx(wb, export_dir / "wrongbook.docx")

            finally:
                st.session_state["job_started"] = False
                st.session_state["phase"] = "done"
                st.rerun()

        out = st.session_state.get("batch_run")
        if out is None:
            if st.session_state["phase"] == "running":
                st.info("批改正在进行中…")
            else:
                st.info("尚未开始或尚未完成批改。")
        else:
            run_dir = Path(out["run_dir"])
            st.success(f"批量完成：{run_dir}")
            st.caption("每个学生一个文件夹，包含 result.json / summary.md / wrongbook.jsonl（若启用）")
            st.json(out)
            
            students = out.get("students", []) or []
            errors = out.get("errors", []) or []

            # Summary table
            rows = []
            for s in students:
                ov = s.get("overall", {})
                if isinstance(ov, dict):
                    v = ov.get("verdict", "")
                    u = bool(ov.get("uncertain", True))
                    overall_str = f"{v}{'（不确定）' if u else ''}"
                else:
                    overall_str = ""
                rows.append(
                    {
                        "student": s.get("student_name", "unknown"),
                        "class": s.get("student_class", ""),
                        "file": s.get("filename", ""),
                        "overall": overall_str,
                        "student_dir": s.get("student_dir", ""),
                    }
                )

            st.dataframe(rows, use_container_width=True, height=320)

            if errors:
                st.warning(f"有 {len(errors)} 个文件处理失败。展开查看错误：")
                with st.expander("错误详情"):
                    st.json(errors)

            st.divider()
            st.markdown("### 按学生查看详情")

            if rows:
                student_names = sorted({r["student"] for r in rows})
                pick = st.selectbox("选择学生", student_names, key="pick_student", disabled=False)

                cand = next((s for s in students if s.get("student_name") == pick), None)
                if cand:
                    student_dir = Path(cand["student_dir"])
                    result_path = student_dir / "result.json"
                    if result_path.exists():
                        res = json.loads(result_path.read_text(encoding="utf-8"))
                        overall = res.get("overall", {})
                        verdict = (
                            overall.get("verdict") if isinstance(overall, dict) else res.get("verdict", "无法判断")
                        ) or "无法判断"
                        uncertain = (
                            bool(overall.get("uncertain", res.get("uncertain", True)))
                            if isinstance(overall, dict)
                            else bool(res.get("uncertain", True))
                        )
                        comment = res.get("comment", "")

                        if uncertain:
                            st.error(f"判定：{verdict}（不确定）")
                        else:
                            st.success(f"判定：{verdict}")

                        if comment:
                            st.write("批注：", comment)

                        items = res.get("items", [])
                        if isinstance(items, list) and items:
                            st.markdown("#### 分题结果")
                            for it in items:
                                if not isinstance(it, dict):
                                    continue
                                qid = it.get("qid", "?")
                                v = it.get("verdict", "无法判断")
                                u = bool(it.get("uncertain", True))
                                cm = it.get("comment", "")

                                st.write(f"**题 {qid}**：{v}{'（不确定）' if u else ''}")
                                if cm:
                                    st.write("评语：", cm)

                                ra = it.get("recognized_answer", "")
                                st.text_area(
                                    "recognized_answer",
                                    value=(ra if isinstance(ra, str) else json.dumps(ra, ensure_ascii=False)) or "（空）",
                                    height=120,
                                    key=f"recognized_answer_{pick}_{qid}",
                                )

                        with st.expander("调试信息（OCR / Spark 原文）"):
                            spark_raw = res.get("spark_raw", "") or ""
                            ocr_raw = res.get("ocr_raw_text", "") or ""

                            # ① 先看长度（这是判断“是否真的被截断”的关键）
                            st.write("len(spark_raw) =", len(spark_raw))

                            # ② Spark 原始输出（完整展示）
                            st.text_area(
                                "Spark raw",
                                value=spark_raw or "（空）",
                                height=240,
                                key=f"spark_raw_{pick}",
                            )

                            # ③ OCR 原始文本
                            st.text_area(
                                "OCR raw",
                                value=ocr_raw or "（空）",
                                height=160,
                                key=f"ocr_raw_{pick}",
                            )
                    else:
                        st.info("该学生目录未找到 result.json")

            st.divider()
            if st.button("导出本次批量错题本（md/docx）", use_container_width=True, key="export_all_wrongbook"):
                if not enable_wrongbook:
                    st.warning("未启用错题本自动保存。请先启用后重新批量运行。")
                else:
                    exported = 0
                    for s in students:
                        student_dir = Path(s.get("student_dir", ""))
                        wb = student_dir / "wrongbook.jsonl"
                        ensure_wrongbook_file(str(wb))
                        export_dir = student_dir / "export"
                        export_dir.mkdir(parents=True, exist_ok=True)

                        export_wrongbook_markdown(str(wb), str(export_dir / "wrongbook.md"))
                        if DOCX_OK:
                            export_wrongbook_docx(wb, export_dir / "wrongbook.docx")
                        exported += 1
                    st.success(f"已为 {exported} 位学生导出错题本到各自 export/ 文件夹。")

    pass


if __name__ == "__main__":
    main()
