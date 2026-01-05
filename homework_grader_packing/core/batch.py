
import io
import json
import hashlib
from datetime import datetime
from pathlib import Path

import pandas as pd
import streamlit as st
from PIL import Image

from .grader_core import (
    AppConfig,
    preprocess_for_ocr,
    batch_grade_images,
    ensure_wrongbook_file,
    export_wrongbook_markdown,
)

# Optional: DOCX export (mirrors app.py behavior)
try:
    from docx import Document
    from docx.shared import Pt
    DOCX_OK = True
except Exception:
    DOCX_OK = False


# -------------------------
# config loader (same spirit as app.py)
# -------------------------
def load_config_from_secrets_or_env() -> AppConfig:
    """
    Load config in this priority:
      1) st.secrets (if present)
      2) config.toml in project root (if present)
      3) OS environment variables
      4) defaults
    Sensitive fields are recommended to be supplied via env/.env, not toml.
    """
    import os

    # Optional: load .env if exists (safe no-op if python-dotenv missing)
    try:
        from dotenv import load_dotenv  # type: ignore
        load_dotenv()  # loads .env from current working dir
    except Exception:
        pass

    # ---------- helpers ----------
    def get_secret(k: str) -> str | None:
        if hasattr(st, "secrets") and k in st.secrets:
            return str(st.secrets[k])
        return None

    def load_toml_dict() -> dict:
        cfg_path = Path("config.toml")
        if not cfg_path.exists():
            return {}
        try:
            # py3.11+: tomllib
            import tomllib  # type: ignore
            return tomllib.loads(cfg_path.read_text(encoding="utf-8"))
        except Exception:
            try:
                import toml  # type: ignore
                return toml.loads(cfg_path.read_text(encoding="utf-8"))
            except Exception:
                return {}

    toml_cfg = load_toml_dict()

    def get_from_toml(k: str, default: str = "") -> str:
        # keys in your toml are flat like PROVIDER, OCR_URL, etc.
        v = toml_cfg.get(k, default)
        return str(v) if v is not None else str(default)

    def get_env(k: str, default: str = "") -> str:
        v = os.environ.get(k, default)
        return str(v) if v is not None else str(default)

    def pick(k: str, default: str = "") -> str:
        # secrets > toml > env > default
        s = get_secret(k)
        if s is not None and s != "":
            return s
        t = get_from_toml(k, "")
        if t != "":
            return t
        e = get_env(k, "")
        if e != "":
            return e
        return default

    # ---------- build cfg ----------
    cfg = AppConfig(
        provider=pick("PROVIDER", "spark_http"),
        spark_http_url=pick("SPARK_HTTP_URL", ""),
        spark_api_password=pick("SPARK_API_PASSWORD", ""),  # SHOULD come from env/.env ideally
        spark_model=pick("SPARK_MODEL", "x1"),
        ocr_url=pick("OCR_URL", ""),
        ocr_app_id=pick("OCR_APP_ID", ""),
        ocr_api_key=pick("OCR_API_KEY", ""),                # SHOULD come from env/.env ideally
        ocr_api_secret=pick("OCR_API_SECRET", ""),          # SHOULD come from env/.env ideally
    )

    return cfg



def _ok_config(cfg: AppConfig) -> bool:
    if cfg.provider == "mock":
        return True
    if cfg.provider in ("cloud_ocr_only", "spark_http"):
        need_ocr = all([cfg.ocr_url, cfg.ocr_app_id, cfg.ocr_api_key, cfg.ocr_api_secret])
        if not need_ocr:
            return False
    if cfg.provider == "spark_http":
        need_spark = bool(cfg.spark_http_url and cfg.spark_api_password)
        if not need_spark:
            return False
    return True


def _export_docx_from_wrongbook(jsonl_path: Path, docx_path: Path) -> int:
    if not DOCX_OK:
        raise RuntimeError("python-docx 未安装，无法导出 docx。")
    rows = []
    if jsonl_path.exists():
        for line in jsonl_path.read_text(encoding="utf-8", errors="ignore").splitlines():
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(line)
                if isinstance(obj, dict):
                    rows.append(obj)
            except Exception:
                pass

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("错题本导出", level=1)
    p = doc.add_paragraph(f"来源：{str(jsonl_path)}\n")
    p.add_run(f"导出时间：{datetime.now().isoformat(timespec='seconds')}\n")
    p.add_run(f"记录条数：{len(rows)}")

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    def qkey(r):
        q = str(r.get("qid", ""))
        try:
            return (0, int(q))
        except Exception:
            return (1, q)

    rows = sorted(rows, key=qkey)
    for r in rows:
        qid = r.get("qid", "")
        subject = r.get("subject", "")
        verdict = r.get("verdict", "")
        uncertain = bool(r.get("uncertain", False))
        rec = (r.get("recognized_answer") or "").strip()
        comment = (r.get("comment") or "").strip()

        doc.add_heading(f"题 {qid}（{subject}）", level=2)
        doc.add_paragraph(f"判定：{verdict}{'（不确定）' if uncertain else ''}")
        if rec:
            doc.add_paragraph(f"学生答案：{rec}")
        if comment:
            doc.add_paragraph(f"评语：{comment}")

    doc.save(str(docx_path))
    return len(rows)


# -------------------------
# UI (match app.py style)
# -------------------------
st.set_page_config(page_title="作业批改助手（批量）", layout="wide")
st.title("作业批改助手（批量版）v0.1（云OCR & Spark）")

cfg0 = load_config_from_secrets_or_env()

with st.sidebar:
    st.header("设置")
    provider = st.selectbox(
        "运行模式",
        ["spark_http", "cloud_ocr_only", "mock"],
        index=["spark_http", "cloud_ocr_only", "mock"].index(cfg0.provider) if cfg0.provider in ["spark_http", "cloud_ocr_only", "mock"] else 0,
        key="provider",
    )

    subject = st.selectbox("科目", ["英语", "科学"], key="subject")
    mode = st.selectbox("批改模式", ["check_and_comment", "check_only"], key="mode")

    st.divider()
    st.subheader("图片预处理")
    crop_ratio = st.slider("中心裁剪比例（越小越聚焦）", min_value=0.60, max_value=1.00, value=1.00, step=0.02, key="crop_ratio")
    max_side = st.selectbox("最大边长(px)", [1200, 1600, 2000], index=1, key="max_side")
    jpeg_quality = st.selectbox("JPEG质量", [60, 70, 75, 80, 85], index=2, key="jpeg_quality")
    enable_cleanup = st.checkbox("用大模型清洗OCR文本", value=True, key="enable_cleanup")

    st.divider()
    st.subheader("题目信息（可选）")
    question_text = st.text_area("题目补充说明（可留空）", value="", height=80, key="question_text")
    answer_key = st.text_input("标准答案（可留空）", value="", key="answer_key")

    st.divider()
    st.subheader("错题本（按学生分开）")
    enable_wrongbook = st.checkbox("自动保存错题记录", value=True, key="enable_wrongbook")
    auto_export_wrongbook = st.checkbox("批量完成后自动导出错题本文档（md/docx）", value=False, key="auto_export_wrongbook")

    st.divider()
    st.subheader("并发（LLM限制=2）")
    llm_concurrency = st.selectbox("LLM并发", [1, 2], index=1, key="llm_concurrency")

# override cfg provider
cfg = AppConfig(
    provider=provider,
    spark_http_url=cfg0.spark_http_url,
    spark_api_password=cfg0.spark_api_password,
    spark_model=cfg0.spark_model,
    ocr_url=cfg0.ocr_url,
    ocr_app_id=cfg0.ocr_app_id,
    ocr_api_key=cfg0.ocr_api_key,
    ocr_api_secret=cfg0.ocr_api_secret,
)

if provider != "mock" and not _ok_config(cfg):
    st.error("配置不完整：请检查 st.secrets 或环境变量中的 OCR_* / SPARK_*。")
    st.stop()

col1, col2 = st.columns([1, 1], gap="large")

with col1:
    st.subheader("上传作业图片（多选）")
    uploaded_files = st.file_uploader(
        "支持 jpg / png（可一次上传 ~50 份）",
        type=["jpg", "jpeg", "png", "webp"],
        accept_multiple_files=True,
        key="batch_uploader",
    )

    if not uploaded_files:
        st.info("请上传多张作业照片或扫描件，然后点击“开始批改”。")
        st.stop()

    st.caption(f"已选择 {len(uploaded_files)} 个文件。")

    # Preview selector (same feel as app.py showing raw + preprocessed)
    names = [f.name for f in uploaded_files]
    preview_name = st.selectbox("预览文件", names, index=0, key="preview_file")
    preview_file = next(f for f in uploaded_files if f.name == preview_name)

    raw_bytes = preview_file.getvalue()
    raw_img = Image.open(io.BytesIO(raw_bytes)).convert("RGB")
    st.image(raw_img, caption="原图（预览）", use_container_width=True)

    pre_bytes = preprocess_for_ocr(
        raw_bytes,
        crop_ratio=crop_ratio,
        max_side=max_side,
        jpeg_quality=jpeg_quality,
        enhance_contrast=1.10,
    )
    pre_img = Image.open(io.BytesIO(pre_bytes)).convert("RGB")
    st.image(pre_img, caption="预处理后（用于OCR，预览）", use_container_width=True)

    st.divider()

    start = st.button("开始批改（批量）", type="primary", use_container_width=True, key="start_batch")

with col2:
    st.subheader("结果")
    if "batch_run" not in st.session_state:
        st.session_state["batch_run"] = None

    if not start and st.session_state["batch_run"] is None:
        st.info("左侧上传图片并点击“开始批改（批量）”。")
        st.stop()

    # run batch
    if start:
        # Build images list
        images = [(f.name, f.getvalue()) for f in uploaded_files]

        with st.spinner("批改中…（批量）"):
            out = batch_grade_images(
                cfg=cfg,
                images=images,
                subject=subject,
                mode=mode,
                question_text=question_text,
                answer_key=answer_key,
                enable_cleanup=enable_cleanup,
                run_root="./runs",
                max_concurrency=int(llm_concurrency),
                crop_ratio=float(crop_ratio),
                max_side=int(max_side),
                jpeg_quality=int(jpeg_quality),
                save_wrongbook=bool(enable_wrongbook),
            )

        st.session_state["batch_run"] = out

        # optional export wrongbook per student
        if enable_wrongbook and auto_export_wrongbook:
            run_dir = Path(out["run_dir"])
            for s in out.get("students", []):
                student_dir = Path(s.get("student_dir", ""))
                wb = student_dir / "wrongbook.jsonl"
                ensure_wrongbook_file(str(wb))
                export_dir = student_dir / "export"
                export_dir.mkdir(parents=True, exist_ok=True)
                export_wrongbook_markdown(str(wb), str(export_dir / "wrongbook.md"))
                if DOCX_OK:
                    _export_docx_from_wrongbook(wb, export_dir / "wrongbook.docx")

    out = st.session_state["batch_run"]
    run_dir = Path(out["run_dir"])

    st.success(f"批量完成：{run_dir}")
    st.caption("每个学生一个文件夹，包含 result.json / summary.md / wrongbook.jsonl（若启用）")

    # Summary table
    students = out.get("students", [])
    errors = out.get("errors", [])

    def overall_verdict(x):
        ov = x.get("overall", {})
        if isinstance(ov, dict):
            v = ov.get("verdict", "")
            u = bool(ov.get("uncertain", True))
            return f"{v}{'（不确定）' if u else ''}"
        return ""

    rows = []
    for s in students:
        rows.append({
            "student": s.get("student_name", "unknown"),
            "class": s.get("student_class", ""),
            "file": s.get("filename", ""),
            "overall": overall_verdict(s),
            "student_dir": s.get("student_dir", ""),
        })

    df = pd.DataFrame(rows).sort_values(["student", "file"]) if rows else pd.DataFrame(columns=["student","class","file","overall","student_dir"])
    st.dataframe(df, use_container_width=True, height=320)

    if errors:
        st.warning(f"有 {len(errors)} 个文件处理失败。展开查看错误：")
        with st.expander("错误详情"):
            st.json(errors)

    st.divider()

    # Per-student drill-down (similar to app.py showing per-question)
    st.markdown("### 按学生查看详情")
    if students:
        pick = st.selectbox("选择学生", [r["student"] for r in rows], key="pick_student")
        cand = next((s for s in students if s.get("student_name") == pick), None)
        if cand:
            student_dir = Path(cand["student_dir"])
            result_path = student_dir / "result.json"
            if result_path.exists():
                res = json.loads(result_path.read_text(encoding="utf-8"))
                overall = res.get("overall", {})
                verdict = (overall.get("verdict") if isinstance(overall, dict) else res.get("verdict", "无法判断")) or "无法判断"
                uncertain = bool(overall.get("uncertain", res.get("uncertain", True))) if isinstance(overall, dict) else bool(res.get("uncertain", True))
                comment = res.get("comment", "")

                if uncertain:
                    st.error(f"判定：{verdict}（不确定）")
                else:
                    st.success(f"判定：{verdict}")

                if comment:
                    st.write("批注：", comment)

                items = res.get("items", [])
                if isinstance(items, list) and items:
                    st.markdown("#### 分题结果")
                    for it in items:
                        if not isinstance(it, dict):
                            continue
                        qid = it.get("qid", "?")
                        v = it.get("verdict", "无法判断")
                        u = bool(it.get("uncertain", True))
                        cm = it.get("comment", "")
                        st.write(f"**题 {qid}**：{v}{'（不确定）' if u else ''}")
                        if cm:
                            st.write("评语：", cm)

                        ra = it.get("recognized_answer", "")
                        # Avoid DuplicateElementId: unique key per qid + student
                        st.text_area(
                            "recognized_answer",
                            value=(ra if isinstance(ra, str) else json.dumps(ra, ensure_ascii=False)) or "（空）",
                            height=120,
                            key=f"recognized_answer_{pick}_{qid}",
                        )

                with st.expander("调试信息（OCR / Spark 原文）"):
                    st.text_area("OCR extracted", value=res.get("ocr_extracted_text", "") or "（空）", height=160, key=f"ocr_ex_{pick}")
                    st.text_area("OCR raw", value=res.get("ocr_raw_text", "") or "（空）", height=160, key=f"ocr_raw_{pick}")
                    st.text_area("Spark raw", value=res.get("spark_raw", "") or "（空）", height=200, key=f"spark_raw_{pick}")

            else:
                st.info("该学生目录未找到 result.json")

    st.divider()

    # Export wrongbook for all students in this run
    c1, c2 = st.columns(2)
    with c1:
        if st.button("导出本次批量错题本（md/docx）", use_container_width=True, key="export_all_wrongbook"):
            if not enable_wrongbook:
                st.warning("未启用错题本自动保存。请先启用后重新批量运行。")
            else:
                exported = 0
                for s in students:
                    student_dir = Path(s.get("student_dir", ""))
                    wb = student_dir / "wrongbook.jsonl"
                    ensure_wrongbook_file(str(wb))
                    export_dir = student_dir / "export"
                    export_dir.mkdir(parents=True, exist_ok=True)
                    export_wrongbook_markdown(str(wb), str(export_dir / "wrongbook.md"))
                    if DOCX_OK:
                        _export_docx_from_wrongbook(wb, export_dir / "wrongbook.docx")
                    exported += 1
                st.success(f"已为 {exported} 位学生导出错题本到各自 export/ 文件夹。")

    with c2:
        idx_path = run_dir / "index.json"
        if idx_path.exists():
            st.download_button(
                "下载 index.json（本次批量清单）",
                data=idx_path.read_bytes(),
                file_name=f"index_{run_dir.name}.json",
                mime="application/json",
                use_container_width=True,
            )
