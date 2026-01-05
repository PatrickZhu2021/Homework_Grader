import base64
import hashlib
import hmac
import io
import json
import re
import time
import os
import json
from datetime import datetime
from collections import defaultdict

# Word 导出用（项目里已有 python-docx 就能用）
from docx import Document
from docx.shared import Pt

from datetime import datetime
from pathlib import Path
from dataclasses import dataclass
from email.utils import formatdate
from typing import Any, Dict, List, Optional
from urllib.parse import quote, urlparse

import requests
import streamlit as st

from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry


from PIL import Image, ImageEnhance


def ensure_str(x):
    if isinstance(x, str):
        return x
    if x is None:
        return ""
    return json.dumps(x, ensure_ascii=False)


def _now_iso() -> str:
    return datetime.now().isoformat(timespec="seconds")


def append_wrongbook_records(
    *,
    result: Dict[str, Any],
    meta: Dict[str, Any],
    out_path: str,
) -> int:
    """Append wrong-question records to a local JSONL file.

    - Writes 1 line per wrong item.
    - Treats verdict != "正确" OR uncertain=True as wrong.
    - Returns number of records written.
    """
    items = result.get("items", [])
    if not isinstance(items, list) or not items:
        return 0

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    written = 0
    with out.open("a", encoding="utf-8") as f:
        for it in items:
            if not isinstance(it, dict):
                continue
            verdict = str(it.get("verdict", "无法判断"))
            uncertain = bool(it.get("uncertain", True))
            if verdict == "正确" and not uncertain:
                continue

            rec = {
                "ts": _now_iso(),
                **meta,
                "qid": str(it.get("qid", "1")),
                "verdict": verdict,
                "uncertain": uncertain,
                "recognized_answer": ensure_str(it.get("recognized_answer", "")),
                "comment": ensure_str(it.get("comment", "")),
            }
            f.write(json.dumps(rec, ensure_ascii=False) + "\n")
            written += 1
    return written

@st.cache_resource
def get_http_session() -> requests.Session:
    s = requests.Session()
    retry = Retry(
        total=4,
        connect=4,
        read=4,
        backoff_factor=0.8,
        status_forcelist=[429, 500, 502, 503, 504],
        allowed_methods=["POST"],
        raise_on_status=False,
    )
    adapter = HTTPAdapter(max_retries=retry, pool_connections=10, pool_maxsize=10)
    s.mount("https://", adapter)
    s.mount("http://", adapter)
    return s



# =========================
# 1) Config
# =========================
@dataclass
class AppConfig:
    provider: str = "mock"  # mock | cloud_ocr_only | spark_http

    # Spark OpenAPI (HTTP)
    spark_http_url: str = ""
    spark_api_password: str = ""
    spark_model: str = "x1"

    # Cloud OCR (OCRforLLM WebAPI)
    ocr_url: str = ""
    ocr_app_id: str = ""
    ocr_api_key: str = ""
    ocr_api_secret: str = ""


def load_config() -> AppConfig:
    cfg = AppConfig()
    cfg.provider = st.secrets.get("PROVIDER", cfg.provider)

    cfg.spark_http_url = st.secrets.get("SPARK_HTTP_URL", cfg.spark_http_url)
    cfg.spark_api_password = st.secrets.get("SPARK_API_PASSWORD", cfg.spark_api_password)
    cfg.spark_model = st.secrets.get("SPARK_MODEL", cfg.spark_model)

    cfg.ocr_url = st.secrets.get("OCR_URL", cfg.ocr_url)
    cfg.ocr_app_id = st.secrets.get("OCR_APP_ID", cfg.ocr_app_id)
    cfg.ocr_api_key = st.secrets.get("OCR_API_KEY", cfg.ocr_api_key)
    cfg.ocr_api_secret = st.secrets.get("OCR_API_SECRET", cfg.ocr_api_secret)
    return cfg


# =========================
# 2) Image preprocess (speed + quality)
# =========================
def preprocess_for_ocr(
    image_bytes: bytes,
    crop_ratio: float = 0.82,
    max_side: int = 1600,
    jpeg_quality: int = 75,
    enhance_contrast: float = 1.10,
) -> bytes:
    """
    Practical OCR preprocessing:
    - center-crop (keeps main content; avoids margins / page numbers)
    - resize so long edge <= max_side
    - mild contrast enhance
    - export to JPEG (smaller upload)
    """
    img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    w, h = img.size

    # center crop
    crop_ratio = max(0.5, min(1.0, crop_ratio))
    cw, ch = int(w * crop_ratio), int(h * crop_ratio)
    left = (w - cw) // 2
    top = (h - ch) // 2
    img = img.crop((left, top, left + cw, top + ch))

    # resize
    w2, h2 = img.size
    scale = min(max_side / max(w2, h2), 1.0)
    if scale < 1.0:
        img = img.resize((int(w2 * scale), int(h2 * scale)))

    # mild contrast
    if enhance_contrast and enhance_contrast != 1.0:
        img = ImageEnhance.Contrast(img).enhance(enhance_contrast)

    out = io.BytesIO()
    img.save(out, format="JPEG", quality=jpeg_quality, optimize=True)
    return out.getvalue()


# =========================
# 3) Cloud OCR (OCRforLLM WebAPI)
# =========================
def _build_ocr_signed_url(ocr_url: str, api_key: str, api_secret: str) -> str:
    u = urlparse(ocr_url)
    host = u.netloc
    path = u.path

    date = formatdate(timeval=None, localtime=False, usegmt=True)
    signature_origin = f"host: {host}\n" f"date: {date}\n" f"POST {path} HTTP/1.1"

    sig_bytes = hmac.new(
        api_secret.encode("utf-8"),
        signature_origin.encode("utf-8"),
        digestmod=hashlib.sha256,
    ).digest()
    signature = base64.b64encode(sig_bytes).decode("utf-8")

    authorization_origin = (
        f'api_key="{api_key}", algorithm="hmac-sha256", '
        f'headers="host date request-line", signature="{signature}"'
    )
    authorization = base64.b64encode(authorization_origin.encode("utf-8")).decode("utf-8")

    return (
        f"{ocr_url}?"
        f"authorization={quote(authorization)}&"
        f"host={quote(host)}&"
        f"date={quote(date)}"
    )


def _safe_json_loads(s: str) -> Optional[Any]:
    try:
        return json.loads(s)
    except Exception:
        return None


def extract_text_from_ocr_obj(obj: Any) -> str:
    """
    Best-effort extraction from OCR JSON structures.
    We try common keys: lines/regions/textline/paragraph/pages/blocks/items, etc.
    """
    if obj is None:
        return ""

    # Direct text
    if isinstance(obj, dict) and isinstance(obj.get("text"), str) and obj["text"].strip():
        return obj["text"].strip()

    lines: List[str] = []

    def walk(o: Any):
        if o is None:
            return
        if isinstance(o, str):
            t = o.strip()
            if t:
                lines.append(t)
            return
        if isinstance(o, dict):
            for k in ["line", "content", "txt", "words", "w", "value"]:
                v = o.get(k)
                if isinstance(v, str) and v.strip():
                    lines.append(v.strip())
            for k in ["lines", "textline", "textlines", "paragraph", "paragraphs", "blocks", "regions", "pages", "items", "data", "result"]:
                if k in o:
                    walk(o[k])
            for v in o.values():
                if isinstance(v, (dict, list)):
                    walk(v)
            return
        if isinstance(o, list):
            for it in o:
                walk(it)

    walk(obj)

    # de-dup
    seen = set()
    out = []
    for t in lines:
        if t in seen:
            continue
        seen.add(t)
        out.append(t)

    return "\n".join(out).strip()


def decode_ocr_payload_text(text_b64: str) -> str:
    if not text_b64:
        return ""
    decoded = base64.b64decode(text_b64).decode("utf-8", errors="ignore").strip()
    if not decoded:
        return ""

    obj = _safe_json_loads(decoded)
    if obj is not None:
        t = extract_text_from_ocr_obj(obj)
        return t if t else decoded
    return decoded


@st.cache_data(show_spinner=False, max_entries=256)
def cloud_ocr_cached(image_sha1: str, image_bytes: bytes, cfg_dict: Dict[str, str]) -> Dict[str, Any]:
    """
    Returns dict: {"raw_text": str, "extracted_text": str, "resp": dict}
    Cached by content hash of preprocessed image bytes.
    """
    ocr_url = cfg_dict["ocr_url"]
    app_id = cfg_dict["ocr_app_id"]
    api_key = cfg_dict["ocr_api_key"]
    api_secret = cfg_dict["ocr_api_secret"]

    signed_url = _build_ocr_signed_url(ocr_url, api_key, api_secret)
    img_b64 = base64.b64encode(image_bytes).decode("utf-8")

    body = {
        "header": {"app_id": app_id, "uid": "hwgrader", "status": 0},
        "parameter": {
            "ocr": {
                "result_option": "normal,char",
                "result_format": "json",
                "output_type": "one_shot",
                "exif_option": "1",
                "alpha_option": "1",
                "rotation_min_angle": 3,
                "result": {"encoding": "utf8", "compress": "raw", "format": "json"},
            }
        },
        "payload": {"image": {"encoding": "jpg", "image": img_b64, "status": 0, "seq": 0}},
    }

    r = requests.post(signed_url, json=body, timeout=35)
    if r.status_code >= 400:
        raise RuntimeError(f"OCR HTTP {r.status_code}: {r.text}")

    data = r.json()
    header = data.get("header", {})
    if header.get("code", 0) != 0:
        raise RuntimeError(f"OCR error: {header}")

    payload = data.get("payload", {}).get("result", {})
    raw_text = decode_ocr_payload_text(payload.get("text", ""))

    obj = _safe_json_loads(raw_text)
    extracted = extract_text_from_ocr_obj(obj) if obj is not None else raw_text

    return {"raw_text": raw_text, "extracted_text": extracted, "resp": data}


def cloud_ocr(image_bytes: bytes, cfg: AppConfig) -> Dict[str, Any]:
    if not (cfg.ocr_url and cfg.ocr_app_id and cfg.ocr_api_key and cfg.ocr_api_secret):
        raise RuntimeError("OCR 配置缺失：请在 secrets.toml 里填 OCR_URL / OCR_APP_ID / OCR_API_KEY / OCR_API_SECRET")

    sha1 = hashlib.sha1(image_bytes).hexdigest()
    cfg_dict = {
        "ocr_url": cfg.ocr_url,
        "ocr_app_id": cfg.ocr_app_id,
        "ocr_api_key": cfg.ocr_api_key,
        "ocr_api_secret": cfg.ocr_api_secret,
    }
    return cloud_ocr_cached(sha1, image_bytes, cfg_dict)


# =========================
# 4) Spark OpenAPI (HTTP)
# =========================
def spark_http_chat(cfg: AppConfig, messages: List[Dict[str, str]], timeout_sec: int = 35, force_json: bool = False) -> str:
    if not cfg.spark_http_url:
        raise RuntimeError("未配置 SPARK_HTTP_URL")
    if not cfg.spark_api_password:
        raise RuntimeError("未配置 SPARK_API_PASSWORD（APIPassword）")

    headers = {"Authorization": f"Bearer {cfg.spark_api_password}", "Content-Type": "application/json"}

    body: Dict[str, Any] = {
        "model": cfg.spark_model,
        "messages": messages,
        "temperature": 0.2,
        "top_k": 4,
        "max_tokens": 900,
        "stream": False,
    }
    if force_json:
        body["response_format"] = {"type": "json_object"}

    session = get_http_session()
    r = session.post(cfg.spark_http_url, headers=headers, json=body, timeout=(10, timeout_sec))
    if r.status_code >= 400:
        raise RuntimeError(f"Spark HTTP {r.status_code}: {r.text}")

    data = r.json()
    try:
        return data["choices"][0]["message"]["content"]
    except Exception:
        return json.dumps(data, ensure_ascii=False)


def try_parse_json(text: str) -> Optional[Dict[str, Any]]:
    """Robust JSON extraction.

    Accepts:
    - pure JSON
    - ```json ... ``` fenced blocks
    - leading/trailing chatter (extracts first JSON object)
    - JSON embedded as a quoted string (e.g. "\"{...}\"")
    """
    if not text:
        return None

    t = text.strip()

    # If the whole payload is a quoted string, decode once (handles escaped newlines, code fences, etc.)
    if (t.startswith('"') and t.endswith('"')) or (t.startswith("'") and t.endswith("'")):
        try:
            decoded = json.loads(t)
            if isinstance(decoded, str):
                t = decoded.strip()
        except Exception:
            pass

    # Strip code fences if present
    if t.startswith("```"):
        t = re.sub(r"^```(?:json)?\s*", "", t, flags=re.IGNORECASE)
        t = re.sub(r"\s*```\s*$", "", t).strip()

    # Fast path
    try:
        obj = json.loads(t)
        return obj if isinstance(obj, dict) else None
    except Exception:
        pass

    # Extract from first '{' to last '}' (best effort)
    i = t.find("{")
    j = t.rfind("}")
    if i == -1 or j == -1 or j <= i:
        return None
    candidate = t[i : j + 1]

    try:
        obj = json.loads(candidate)
        return obj if isinstance(obj, dict) else None
    except Exception:
        return None


def pick_exam_region(text: str, max_chars: int = 2200) -> str:
    """Keep mostly the question/answer region to reduce LLM tokens and latency."""
    if not text:
        return ""
    t = text

    for marker in ["[基础过关训练]", "基础过关训练", "试题", "练习", "一、", "一."]:
        idx = t.find(marker)
        if idx != -1:
            t = t[idx:]
            break

    t = t.strip()
    if len(t) > max_chars:
        t = t[:max_chars] + "\n(…后略…)"
    return t

def llm_cleanup_ocr(cfg: AppConfig, extracted_text: str) -> str:
    extracted_text = (extracted_text or "").strip()
    if not extracted_text:
        return ""

    prompt = f"""
你是OCR文本清洗器。给你一段OCR识别的文本，请你做“轻度清洗整理”：
- 不要编造未出现的信息
- 删除明显无关的页码/噪声（比如孤立的编号、重复）
- 合并被错误断行的句子
- 保留数学符号、单位、数字
只输出清洗后的文本本身，不要输出任何解释或多余文字。

OCR文本：
{extracted_text}
""".strip()

    out = spark_http_chat(cfg, [{"role": "user", "content": prompt}], timeout_sec=25, force_json=False)
    return out.strip()

def build_grading_prompt(subject, mode, question_text, answer_key, student_text, enable_cleanup: bool = True) -> str:
    '''
    Minimal JSON prompt: the model outputs ONLY {"overall":..., "items":[...]}.

    We DO NOT ask the model to echo OCR text (raw/clean/struct) to avoid token blowups and truncated JSON.
    OCR/debug fields are attached locally by Python.
    '''
    comment_req = "每道题 comment 写 1-2 句话（各题独立）。" if mode == "check_and_comment" else "每道题 comment 写 1 句话以内或写“OK”。"
    cleanup_req = (
        "你可以做轻度清洗：纠正明显OCR错字/断行/空格，但不要凭空补充内容。"
        if enable_cleanup else
        "不要做任何清洗或纠错：严格基于原始OCR文本判题。"
    )

    # IMPORTANT: this is an f-string; any literal braces must be doubled.
    schema_example = (
        "{{\n"
        '  "overall": {{"verdict": "正确", "uncertain": false}},\n'
        '  "items": [\n'
        '    {{"qid": "1", "verdict": "正确", "uncertain": false, "recognized_answer": "…", "comment": "…"}}\n'
        "  ]\n"
        "}}"
    )

    return f"""
你是一位严格的作业批改助手。你会收到 OCR 识别到的学生作答文本（可能有错字/断行）、题目文本、标准答案/评分要点。
你必须只输出严格 JSON（不要 Markdown、不要代码块、不要解释性文字）。

科目：{subject}
题目/说明（可能包含多题）：{(question_text or "").strip() or "无"}
标准答案（可能包含多题）：{(answer_key or "").strip() or "无"}

学生作答文本（OCR 原文）：
{(student_text or "").strip() or "（未识别到学生答案）"}

规则：
- {cleanup_req}
- 对每道题给出 verdict：只能是 正确/错误/接近正确/部分正确/无法判断
- {comment_req}
- 如果题目无法明确拆分多题，也必须返回 items 且至少包含一个元素 qid="1"。

输出 JSON schema（必须完全匹配，且只输出这一份 JSON）：
{schema_example}
""".strip()


# =========================
# 5) Providers
# =========================
def mock_provider() -> Dict[str, Any]:
    return {"verdict": "无法判断", "comment": "mock", "uncertain": True, "recognized_answer": ""}


def cloud_ocr_only_provider(cfg: AppConfig, pre_bytes: bytes) -> Dict[str, Any]:
    t0 = time.time()
    ocr = cloud_ocr(pre_bytes, cfg)
    dt = time.time() - t0
    text = (ocr.get("extracted_text") or "").strip()
    return {"verdict": "无法判断", "comment": f"OK（OCR {dt:.1f}s）", "uncertain": (not bool(text)), "recognized_answer": text}

def spark_http_provider(
    cfg: AppConfig,
    pre_bytes: bytes,
    subject: str,
    mode: str,
    question_text: str,
    answer_key: str,
    enable_cleanup: bool,
) -> Dict[str, Any]:
    """
    Provider: Cloud OCR -> (optional) LLM cleanup -> Spark grading

    Guarantees:
    - Always returns OCR extracted text, OCR raw text, OCR full response, Spark raw output (if any)
    - Retries Spark call on transient failures
    - Supports multi-question output (overall/items) if prompt asks for it
      and provides backward-compatible top-level verdict/comment/uncertain/recognized_answer.
    """
    # ---------- OCR ----------
    t0 = time.time()
    ocr = cloud_ocr(pre_bytes, cfg)  # dict: raw_text / extracted_text / resp
    ocr_dt = time.time() - t0

    image_sha1 = hashlib.sha1(pre_bytes).hexdigest()
    extracted = (ocr.get("extracted_text") or "").strip()

    # Keep only the most relevant region for LLM to reduce latency
    extracted_region = pick_exam_region(extracted, max_chars=2200)

    cleaned = extracted_region
    cleanup_dt = 0.0

    # ---------- Optional cleanup (LLM) ----------
    if enable_cleanup and extracted_region:
        t_clean = time.time()
        try:
            cleaned = llm_cleanup_ocr(cfg, extracted_region[:1800]).strip()
        except Exception:
            cleaned = extracted_region
        cleanup_dt = time.time() - t_clean

    # ---------- Build prompt ----------
    prompt = build_grading_prompt(subject, mode, question_text, answer_key, cleaned, enable_cleanup=False)

    # Common debug payload we always include
    base_debug = {
        "image_sha1": image_sha1,
        "ocr_extracted_text": extracted,
        "ocr_clean_text": cleaned,
        "ocr_raw_text": ocr.get("raw_text", ""),
        "ocr_response": ocr.get("resp", {}),
        "spark_raw": "",
    }

    # ---------- Spark call with retries ----------
    last_err: Optional[Exception] = None
    spark_dt = 0.0
    for attempt in range(3):  # 3 tries is a good balance
        try:
            t_spark = time.time()
            out = spark_http_chat(
                cfg,
                [{"role": "user", "content": prompt}],
                timeout_sec=60,      # allow slow responses
                force_json=True,     # we want JSON when possible
            )
            spark_dt = time.time() - t_spark
            base_debug["spark_raw"] = out

            parsed = try_parse_json(out)
            if not parsed:
                # Retry once with shorter input + stricter prompt (helps first-call format drift)
                retry_prompt = build_grading_prompt(
                    subject,
                    mode,
                    question_text,
                    answer_key,
                    pick_exam_region(cleaned or extracted_region or extracted, max_chars=1200),
                    enable_cleanup=False,
                )
                try:
                    t_retry = time.time()
                    out2 = spark_http_chat(
            cfg,
            [{"role": "user", "content": retry_prompt}],
            timeout_sec=45,
            force_json=True,
                    )
                    spark_dt = spark_dt + (time.time() - t_retry)
                    base_debug["spark_raw"] = out2
                    parsed = try_parse_json(out2)
                except Exception:
                    parsed = None

                if not parsed:
                    timing = f"OCR {ocr_dt:.1f}s, LLM {spark_dt:.1f}s"
                    return {
            "verdict": "无法判断",
            "comment": f"模型输出不是JSON（{timing}）",
            "uncertain": True,
            "recognized_answer": cleaned or extracted_region or extracted,
            # Multi-question fields (empty but present)
            "overall": {"verdict": "无法判断", "uncertain": True},
            "items": [],
            **base_debug,
                    }



            # ---- Support BOTH schemas ----
            # A) New schema: {"overall": {...}, "items": [...]}
            items = parsed.get("items", None)
            overall = parsed.get("overall", None)

            # Validate items schema
            if isinstance(items, list) and (overall is None or isinstance(overall, dict)):
                # Normalize overall
                if not isinstance(overall, dict):
                    overall = {}
                o_verdict = overall.get("verdict", "无法判断")
                o_uncertain = bool(overall.get("uncertain", True))

                allowed = {"正确", "错误", "接近正确", "部分正确", "无法判断"}
                if o_verdict not in allowed:
                    o_verdict = "无法判断"
                    o_uncertain = True

                # Normalize each item
                norm_items = []
                for it in items:
                    if not isinstance(it, dict):
                        continue
                    qid = str(it.get("qid", "") or "1")
                    v = it.get("verdict", "无法判断")
                    u = bool(it.get("uncertain", True))
                    ra = it.get("recognized_answer", "")
                    cm = it.get("comment", "")

                    if v not in allowed:
                        v = "无法判断"
                        u = True

                    # Ensure comment length / presence in comment mode
                    if mode == "check_and_comment":
                        if not isinstance(cm, str) or len(cm.strip()) == 0:
                            cm = "答案要点未识别清楚，请重拍更清晰。"
                        # keep it short-ish, but 1-2 sentences
                        cm = cm.strip()

                    norm_items.append(
                        {
                            "qid": qid,
                            "verdict": v,
                            "uncertain": u,
                            "recognized_answer": ra if isinstance(ra, str) else json.dumps(ra, ensure_ascii=False),
                            "comment": cm if isinstance(cm, str) else json.dumps(cm, ensure_ascii=False),
                        }
                    )

                timing = f"OCR {ocr_dt:.1f}s, LLM {spark_dt:.1f}s"
                # Backward-compatible top-level fields
                recognized_raw = ensure_str(parsed.get('recognized_answer_raw', extracted))
                recognized_clean = ensure_str(parsed.get('recognized_answer_clean', extracted))
                recognized_struct = parsed.get('recognized_answer_struct', {})
                base_debug["ocr_clean_text"] = recognized_clean
                top_recognized = recognized_clean or recognized_raw or extracted
                return {
                    "overall": {"verdict": o_verdict, "uncertain": o_uncertain},
                    "items": norm_items,
                    "recognized_answer_raw": recognized_raw,
                    "recognized_answer_clean": recognized_clean,
                    "recognized_answer_struct": recognized_struct,
                    "verdict": o_verdict,
                    "uncertain": o_uncertain,
                    "comment": f"OK（{timing}）",
                    "recognized_answer": top_recognized,
                    **base_debug,
                }

            # B) Old schema: {"verdict":..., "comment":..., "uncertain":..., "recognized_answer":...}
            verdict = parsed.get("verdict", "无法判断")
            comment = parsed.get("comment", "")
            uncertain = bool(parsed.get("uncertain", True))
            recognized_answer = parsed.get("recognized_answer", "") or (cleaned or extracted)

            allowed = {"正确", "错误", "接近正确", "部分正确", "无法判断"}
            if verdict not in allowed:
                verdict = "无法判断"
                uncertain = True

            # If student text too short -> force uncertain
            if len((cleaned or "").strip()) < 12:
                uncertain = True
                if not comment:
                    comment = "答案不清晰"

            # Ensure comment mode: 1-2 sentences even for single question
            if mode == "check_and_comment":
                if not isinstance(comment, str) or len(comment.strip()) == 0:
                    comment = "答案要点未识别清楚，请重拍更清晰。"
                comment = comment.strip()

            timing = f"OCR {ocr_dt:.1f}s, LLM {spark_dt:.1f}s"
            comment = (comment + f"（{timing}）") if isinstance(comment, str) and comment else f"OK（{timing}）"

            # Also provide items[] even in old schema (single item) for uniform UI
            single_item = {
                "qid": "1",
                "verdict": verdict,
                "uncertain": uncertain,
                "recognized_answer": recognized_answer if isinstance(recognized_answer, str) else json.dumps(recognized_answer, ensure_ascii=False),
                "comment": comment,
            }

            return {
                "overall": {"verdict": verdict, "uncertain": uncertain},
                "items": [single_item],
                "verdict": verdict,
                "comment": comment,
                "uncertain": uncertain,
                "recognized_answer": recognized_answer if isinstance(recognized_answer, str) else json.dumps(recognized_answer, ensure_ascii=False),
                **base_debug,
            }

        except Exception as e:
            last_err = e
            # exponential-ish backoff
            time.sleep(1.0 * (attempt + 1))
            continue

    # All retries failed: return graceful fallback with OCR details preserved
    timing = f"OCR {ocr_dt:.1f}s, LLM {spark_dt:.1f}s"
    return {
        "verdict": "无法判断",
        "comment": f"Spark异常（{timing}）：{last_err}",
        "uncertain": True,
        "recognized_answer": extracted,
        "overall": {"verdict": "无法判断", "uncertain": True},
        "items": [],
        **base_debug,
    }

def _safe_read_jsonl(jsonl_path: str):
    rows = []
    if not jsonl_path or not os.path.exists(jsonl_path):
        return rows
    with open(jsonl_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                rows.append(json.loads(line))
            except Exception:
                # 单行坏了就跳过，不影响导出
                continue
    return rows


def _norm_date(ts: str) -> str:
    """把时间戳归一成 YYYY-MM-DD；ts 为空就返回 unknown"""
    if not ts:
        return "unknown"
    # ts 可能是 2025-12-29T10:52:12 或带时区等，取前 10 位通常够用
    if len(ts) >= 10 and ts[4] == "-" and ts[7] == "-":
        return ts[:10]
    return "unknown"


def _student_key(row: dict) -> str:
    clazz = (row.get("student_class") or "").strip()
    name = (row.get("student_name") or "").strip()
    if clazz and name:
        return f"{clazz} - {name}"
    if name:
        return name
    if clazz:
        return clazz
    return "未填写学生信息"

def ensure_wrongbook_file(path: str) -> None:
    """确保 wrongbook 的目录和 jsonl 文件存在（即使还没有任何记录）。"""
    if not path:
        return
    folder = os.path.dirname(path) or "."
    os.makedirs(folder, exist_ok=True)
    if not os.path.exists(path):
        # 创建空文件（也可以写一行 meta/header，看你要不要）
        with open(path, "w", encoding="utf-8") as f:
            # 空文件即可；若你希望首行写入导出信息，可取消注释：
            # f.write("")  
            pass


def export_wrongbook_documents(jsonl_path: str, out_dir: str, base_name: str = "wrongbook_export"):
    """
    从 JSONL 导出 Markdown + DOCX。
    返回 (md_path, docx_path, count)
    """
    rows = _safe_read_jsonl(jsonl_path)
    if not rows:
        return None, None, 0

    os.makedirs(out_dir, exist_ok=True)
    md_path = os.path.join(out_dir, f"{base_name}.md")
    docx_path = os.path.join(out_dir, f"{base_name}.docx")

    # 结构：student -> date -> list[rows]
    grouped = defaultdict(lambda: defaultdict(list))
    for r in rows:
        grouped[_student_key(r)][_norm_date(r.get("ts"))].append(r)

    # ---- Markdown ----
    def _qid_sort_key(x):
        q = str(x.get("qid", ""))
        # 尝试把 "3" 这种排序到前面；非数字放后面
        try:
            return (0, int(q))
        except Exception:
            return (1, q)

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(f"# 错题本导出\n\n")
        f.write(f"- 来源：{jsonl_path}\n")
        f.write(f"- 导出时间：{datetime.now().isoformat(timespec='seconds')}\n")
        f.write(f"- 记录条数：{len(rows)}\n\n")

        for student in sorted(grouped.keys()):
            f.write(f"## {student}\n\n")
            dates = grouped[student]
            for d in sorted(dates.keys()):
                f.write(f"### {d}\n\n")
                items = sorted(dates[d], key=_qid_sort_key)
                for r in items:
                    qid = r.get("qid", "")
                    verdict = r.get("verdict", "")
                    uncertain = r.get("uncertain", False)
                    rec = (r.get("recognized_answer") or "").strip()
                    comment = (r.get("comment") or "").strip()
                    subject = (r.get("subject") or "").strip()

                    f.write(f"#### 题 {qid}（{subject}）\n\n")
                    f.write(f"- 判定：{verdict}{'（不确定）' if uncertain else ''}\n")
                    if rec:
                        f.write(f"- 学生答案：{rec}\n")
                    if comment:
                        f.write(f"- 评语：{comment}\n")
                    f.write("\n")

    # ---- DOCX ----
    doc = Document()
    doc.add_heading("错题本导出", level=1)
    p = doc.add_paragraph(f"来源：{jsonl_path}\n")
    p.add_run(f"导出时间：{datetime.now().isoformat(timespec='seconds')}\n")
    p.add_run(f"记录条数：{len(rows)}")

    # 统一字体稍微好看点（可选）
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    for student in sorted(grouped.keys()):
        doc.add_heading(student, level=2)
        dates = grouped[student]
        for d in sorted(dates.keys()):
            doc.add_heading(d, level=3)
            items = sorted(dates[d], key=_qid_sort_key)
            for r in items:
                qid = r.get("qid", "")
                verdict = r.get("verdict", "")
                uncertain = r.get("uncertain", False)
                rec = (r.get("recognized_answer") or "").strip()
                comment = (r.get("comment") or "").strip()
                subject = (r.get("subject") or "").strip()

                doc.add_heading(f"题 {qid}（{subject}）", level=4)
                doc.add_paragraph(f"判定：{verdict}{'（不确定）' if uncertain else ''}")
                if rec:
                    doc.add_paragraph(f"学生答案：{rec}")
                if comment:
                    doc.add_paragraph(f"评语：{comment}")

    doc.save(docx_path)

    return md_path, docx_path, len(rows)



# =========================
# 6) Streamlit UI
# =========================
st.set_page_config(page_title="作业批改助手", layout="wide")
st.title("作业批改助手v0.4.0（云OCR & Spark）")

cfg = load_config()

with st.sidebar:
    st.header("设置")
    provider = st.selectbox(
        "运行模式",
        ["spark_http", "cloud_ocr_only", "mock"],
        index=["spark_http", "cloud_ocr_only", "mock"].index(cfg.provider) if cfg.provider in ["spark_http", "cloud_ocr_only", "mock"] else 0,
    )

    subject = st.selectbox("科目", ["英语", "科学"])
    mode = st.selectbox("批改模式", [ "check_and_comment", "check_only"])

    st.divider()
    st.subheader("图片预处理")
    crop_ratio = st.slider("中心裁剪比例（越小越聚焦）", min_value=0.60, max_value=1.00, value=1.00, step=0.02)
    max_side = st.selectbox("最大边长(px)", [1200, 1600, 2000], index=1)
    jpeg_quality = st.selectbox("JPEG质量", [60, 70, 75, 80, 85], index=2)
    enable_cleanup = st.checkbox("用大模型清洗OCR文本", value=True)

    st.divider()
    st.subheader("题目信息（可选）")
    question_text = st.text_area("题目补充说明（可留空）", value="", height=80)
    answer_key = st.text_input("标准答案（可留空）", value="")

    st.divider()
    st.subheader("错题本（本地自动保存）")
    enable_wrongbook = st.checkbox("自动保存错题记录", value=False)
    student_name = st.text_input("学生姓名/编号（可留空）", value="")
    student_class = st.text_input("班级（可留空）", value="")
    wrongbook_path = st.text_input("保存路径（JSONL）", value="./wrongbook/wrongbook.jsonl")

col1, col2 = st.columns([1, 1], gap="large")

with col1:
    st.subheader("上传作业图片")
    uploaded = st.file_uploader("支持 jpg / png", type=["jpg", "jpeg", "png"])
    if not uploaded:
        st.info("请先上传一张作业照片或扫描件。")
        st.stop()

    raw_bytes = uploaded.getvalue()
    raw_img = Image.open(io.BytesIO(raw_bytes)).convert("RGB")
    st.image(raw_img, caption="原图", use_container_width=True)

    pre_bytes = preprocess_for_ocr(
        raw_bytes,
        crop_ratio=crop_ratio,
        max_side=max_side,
        jpeg_quality=jpeg_quality,
        enhance_contrast=1.10,
    )
    pre_img = Image.open(io.BytesIO(pre_bytes)).convert("RGB")
    st.image(pre_img, caption="预处理后（用于OCR）", use_container_width=True)

    st.divider()
    if st.button("开始批改", type="primary", use_container_width=True):
        if enable_wrongbook:
            ensure_wrongbook_file(wrongbook_path)
        with st.spinner("批改中..."):
            try:
                if provider == "mock":
                    result = mock_provider()
                elif provider == "cloud_ocr_only":
                    result = cloud_ocr_only_provider(cfg, pre_bytes)
                else:
                    result = spark_http_provider(
                        cfg,
                        pre_bytes,
                        subject,
                        mode,
                        question_text,
                        answer_key,
                        enable_cleanup,
                    )
            except Exception as e:
                # 保底：不要丢掉UI；尽量保留OCR文本（如果 provider 没跑到就为空）
                result = {
                    "verdict": "无法判断",
                    "comment": f"运行出错：{e}",
                    "uncertain": True,
                    "recognized_answer": "",
                    "overall": {"verdict": "无法判断", "uncertain": True},
                    "items": [],
                    "ocr_extracted_text": "",
                    "ocr_raw_text": "",
                    "ocr_response": {},
                    "spark_raw": "",
                }

        # --- Auto-save wrongbook (local) ---
        if enable_wrongbook and isinstance(result, dict):
            try:
                image_sha1 = str(result.get("image_sha1", ""))
                # Avoid duplicate appends when user re-clicks the button on the same image/result
                save_key = f"{image_sha1}:{provider}:{subject}:{mode}:{hashlib.md5(json.dumps(result.get('items', []), ensure_ascii=False).encode('utf-8')).hexdigest()}"
                if st.session_state.get("last_wrongbook_save_key") != save_key:
                    meta = {
                        "provider": provider,
                        "subject": subject,
                        "mode": mode,
                        "student_name": student_name.strip(),
                        "student_class": student_class.strip(),
                        "image_sha1": image_sha1,
                        "filename": getattr(uploaded, "name", ""),
                    }
                    written = append_wrongbook_records(result=result, meta=meta, out_path=wrongbook_path)
                    st.session_state["last_wrongbook_save_key"] = save_key
                    if written > 0:
                        st.toast(f"已保存 {written} 条错题记录到 {wrongbook_path}")
            except Exception as e:
                st.warning(f"错题本保存失败：{e}")

        st.session_state["last_result"] = result

with col2:
    st.subheader("结果")
    if "last_result" not in st.session_state:
        st.info("上传图片并点击“开始批改”。")
        st.stop()

    res = st.session_state["last_result"]

    # ---- Overall verdict ----
    overall = res.get("overall")
    if isinstance(overall, dict):
        verdict = overall.get("verdict", res.get("verdict", "无法判断"))
        uncertain = bool(overall.get("uncertain", res.get("uncertain", True)))
    else:
        verdict = res.get("verdict", "无法判断")
        uncertain = bool(res.get("uncertain", True))

    comment = res.get("comment", "")

    if uncertain:
        st.error(f"判定：{verdict}（不确定）")
    else:
        st.success(f"判定：{verdict}")

    if comment:
        st.write("批注：", comment)

    # ---- Per-question items (comment mode wants 1-2 sentences each) ----
    items = res.get("items", [])
    if isinstance(items, list) and len(items) > 0:
        st.markdown("### 分题结果")
        for it in items:
            if not isinstance(it, dict):
                continue
            qid = it.get("qid", "?")
            v = it.get("verdict", "无法判断")
            u = bool(it.get("uncertain", True))
            cm = it.get("comment", "")
            ra = it.get("recognized_answer", "")

            badge = "（不确定）" if u else ""
            st.markdown(f"**题 {qid}：{v}{badge}**")
            if cm:
                st.write(cm)
            with st.expander(f"题 {qid} 识别到的答案"):
                st.text_area("recognized_answer", value=ensure_str(ra) or "（空）", height=120, key=f"recognized_answer_{qid}")

    # ---- Always show OCR/cleaned text used for grading ----
    ocr_used = ensure_str(res.get("recognized_answer_clean", res.get("ocr_clean_text", res.get("ocr_extracted_text", res.get("recognized_answer", "")))))
    st.divider()
    st.text_area(
        "识别/清洗后的文本（给大模型判题用）",
        value=ocr_used or "（空）",
        height=320,
    )
    if len(ocr_used.strip()) < 12:
        st.warning("识别到的答案很短。建议：裁剪比例调小一点（如 0.78），或拍近一点/更清晰。")

    with st.expander("OCR 原始信息"):
        st.text_area("OCR extracted（给模型用，原文）", value=ensure_str(res.get("ocr_extracted_text", "")), height=180)
        st.text_area("LLM 清洗后文本（若启用/可用）", value=ensure_str(res.get("recognized_answer_clean", res.get("ocr_clean_text", ""))), height=180)
        st.text_area("OCR raw_text（payload解码）", value=ensure_str(res.get("ocr_raw_text", "")), height=180)
        st.json(res.get("ocr_response", {}))

    with st.expander("Spark 原始输出"):
        st.text_area("spark_raw", value=ensure_str(res.get("spark_raw", "")), height=220)

    with st.expander("完整结果 JSON"):
        st.code(json.dumps(res, ensure_ascii=False, indent=2), language="json")

st.sidebar.markdown("---")
st.sidebar.subheader("错题本导出")

export_out_dir = st.sidebar.text_input(
    "导出目录",
    value=os.path.join(".", "wrongbook"),
    help="导出 Markdown 和 Word 的目录"
)

export_base_name = st.sidebar.text_input(
    "导出文件名（不带后缀）",
    value="wrongbook_export"
)

if st.sidebar.button("导出错题本文档", use_container_width=True):
    try:
        ensure_wrongbook_file(wrongbook_path)  # ✅ 保证文件/目录存在

        md_path, docx_path, n = export_wrongbook_documents(
            jsonl_path=wrongbook_path,
            out_dir=export_out_dir,
            base_name=export_base_name
        )

        if n == 0:
            st.sidebar.info("错题本目前为空（已创建 JSONL 文件），没有可导出的内容。")
        else:
            st.sidebar.success(f"导出完成：{n} 条记录")
            st.sidebar.write(f"- Markdown：{md_path}")
            st.sidebar.write(f"- Word：{docx_path}")
    except Exception as e:
        st.sidebar.error(f"导出失败：{e}")
